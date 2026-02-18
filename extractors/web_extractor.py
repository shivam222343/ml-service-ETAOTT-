import requests
from bs4 import BeautifulSoup
import os
import json
import cloudinary.uploader
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import tempfile
from dotenv import load_dotenv
import html2text
import re
from collections import Counter

load_dotenv()

# Configure Cloudinary
cloudinary.config(
    cloud_name=os.getenv('CLOUDINARY_CLOUD_NAME'),
    api_key=os.getenv('CLOUDINARY_API_KEY'),
    api_secret=os.getenv('CLOUDINARY_API_SECRET'),
    secure=True
)

def extract_web_content(url):
    """
    Extract content from a web page, simplify it using AI, and generate a thumbnail.
    """
    try:
        print(f"🌐 Scraping web content from: {url}")
        
        # 1. Fetch and Scrape Content
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
            script.decompose()
            
        # Get title
        title = soup.title.string if soup.title else "Web Resource"
        
        # Convert HTML to Markdown for better processing
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        h.body_width = 0
        raw_text = h.handle(response.text)
        
        # Clean up text
        cleaned_text = re.sub(r'\n\s*\n', '\n\n', raw_text).strip()
        
        # 2. AI Simplification (using Groq)
        simplified_text = cleaned_text
        summary = ""
        topics = []
        keywords = []
        
        groq_api_key = os.getenv('GROQ_API_KEY')
        if groq_api_key:
            try:
                print("🧠 Simplifying content with AI...")
                # Prepare prompt for student-level conversion
                prompt = f"""
                You are a world-class academic tutor. I will provide you with raw text scraped from a website. 
                Your task is to:
                1. Convert the content into a student-level readable and highly understandable format.
                2. Structure it with clear headings, bullet points, and simple explanations.
                3. Extract a short 2-3 sentence summary.
                4. List 3-5 main topics.
                5. List 5-10 key terms.

                Return the result in the following JSON format ONLY:
                {{
                    "simplified_content": "MARGON_CONTENT_HERE",
                    "summary": "SUMMARY_HERE",
                    "topics": ["topic1", "topic2"],
                    "keywords": ["key1", "key2"]
                }}

                RAW TEXT:
                {cleaned_text[:4000]} 
                """
                
                groq_response = requests.post(
                    'https://api.groq.com/openai/v1/chat/completions',
                    headers={
                        'Authorization': f'Bearer {groq_api_key}',
                        'Content-Type': 'application/json'
                    },
                    json={
                        'model': 'llama-3.3-70b-versatile',
                        'messages': [{'role': 'user', 'content': prompt}],
                        'temperature': 0.3
                    }
                )
                
                if groq_response.status_code == 200:
                    ai_data = groq_response.json()
                    content_str = ai_data['choices'][0]['message']['content']
                    
                    # Try to parse JSON from AI response
                    try:
                        # Extract JSON if wrapped in backticks
                        json_match = re.search(r'\{[\s\S]*\}', content_str)
                        if json_match:
                            parsed_ai = json.loads(json_match.group())
                            simplified_text = parsed_ai.get('simplified_content', cleaned_text)
                            summary = parsed_ai.get('summary', "")
                            topics = parsed_ai.get('topics', [])
                            keywords = parsed_ai.get('keywords', [])
                        else:
                            print("⚠️ AI did not return valid JSON, using fallback parsing")
                    except Exception as e:
                        print(f"⚠️ JSON Parse error for AI response: {e}")
            except Exception as e:
                print(f"⚠️ AI Simplification failed: {e}")

        # Fallback for topics/keywords if AI failed or not available
        if not topics:
            words = re.findall(r'\w{5,}', cleaned_text.lower())
            common_words = {'about', 'after', 'again', 'could', 'every', 'from', 'great', 'have', 'their', 'there', 'these', 'which', 'would'}
            keywords = [word for word, count in Counter(words).most_common(10) if word not in common_words]
            topics = [word.capitalize() for word, count in Counter(words).most_common(3) if word not in common_words]
            summary = cleaned_text[:500] + "..."

        # 3. Generate Thumbnail (Screenshot)
        thumbnail_url = None
        thumbnail_public_id = None
        
        thumbnail_bytes = None
        try:
            print("📸 Attempting to generate page screenshot...")
            # We'll use a screenshot API or playwright if possible.
            # For this environment, since we can't easily install playwright browsers mid-run,
            # we'll use a reliable screenshot service or a generic web thumbnail if it fails.
            
            # Trying Playwright if installed
            try:
                from playwright.sync_api import sync_playwright
                print("🎮 Attempting Playwright screenshot...")
                with sync_playwright() as p:
                    browser = p.chromium.launch(headless=True)
                    context = browser.new_context(
                        user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                        viewport={'width': 1280, 'height': 720}
                    )
                    page = context.new_page()
                    page.goto(url, wait_until="domcontentloaded", timeout=30000)
                    thumbnail_bytes = page.screenshot(type='jpeg', quality=80)
                    browser.close()
                    
                    # Upload to Cloudinary
                    upload_result = cloudinary.uploader.upload(
                        thumbnail_bytes,
                        folder="eta-web-thumbnails",
                        resource_type="image"
                    )
                    thumbnail_url = upload_result.get("secure_url")
                    thumbnail_public_id = upload_result.get("public_id")
                    print(f"✅ Web Screenshot uploaded: {thumbnail_url}")
            except Exception as pw_err:
                print(f"⚠️ Playwright screenshot skipped: {pw_err}")
                print("ℹ️ Note: Screenshots require system-level browser dependencies not available in this environment.")
                # Fallback: Use a generic web icon or a simple card-style image
                # Actually, let's use a nice looking placeholder from a UI library or just skip
        except Exception as thumb_err:
            print(f"⚠️ Web Thumbnail generation failed: {thumb_err}")

        # 4. Generate PDF version
        pdf_url = None
        pdf_public_id = None
        try:
            print("📄 Generating PDF summary...")
            pdf = FPDF()
            pdf.add_page()
            
            # Title
            pdf.set_font("Helvetica", 'B', 16)
            pdf.cell(0, 10, title, ln=True, align='C')
            pdf.line(10, 20, 200, 20)
            pdf.ln(5)
            
            # Add Screenshot if available
            screenshot_temp_path = None
            if thumbnail_bytes:
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
                        tmp.write(thumbnail_bytes)
                        screenshot_temp_path = tmp.name
                    pdf.image(screenshot_temp_path, x=45, y=30, w=120)
                    pdf.ln(75)
                except Exception as img_err:
                    print(f"⚠️ Could not add image to PDF: {img_err}")
            else:
                pdf.ln(10)
            
            # Summary Section
            pdf.set_font("Helvetica", 'B', 12)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 10, "Executive Summary", ln=True, fill=True)
            pdf.set_font("Helvetica", '', 10)
            pdf.multi_cell(0, 7, summary.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(5)
            
            # Topics
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 10, "Key Insights & Topics", ln=True)
            pdf.set_font("Helvetica", '', 10)
            topics_str = " - " + " - ".join(topics)
            pdf.cell(0, 10, topics_str.encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.ln(5)
            
            # Simplified Content
            pdf.set_font("Helvetica", 'B', 12)
            pdf.cell(0, 10, "Simplified Explanation", ln=True)
            pdf.set_font("Helvetica", '', 10)
            clean_content = simplified_text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, clean_content)
            
            # Save to temporary file and upload
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                pdf.output(tmp_pdf.name)
                print("☁️ Uploading PDF to Cloudinary...")
                pdf_upload = cloudinary.uploader.upload(
                    tmp_pdf.name,
                    folder="eta-web-pdfs",
                    resource_type="raw"
                )
                pdf_url = pdf_upload.get("secure_url")
                pdf_public_id = pdf_upload.get("public_id")
                os.unlink(tmp_pdf.name)
            
            print(f"✅ Web PDF version uploaded: {pdf_url}")
        except Exception as pdf_err:
            print(f"⚠️ PDF generation failed: {pdf_err}")

        # 5. Generate Word version
        docx_url = None
        docx_public_id = None
        try:
            print("📝 Generating Word document...")
            doc = Document()
            doc.add_heading(title, 0)
            
            if screenshot_temp_path:
                try:
                    doc.add_picture(screenshot_temp_path, width=Inches(5))
                except: pass
            
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(summary)
            
            doc.add_heading('Key Insights', level=1)
            p = doc.add_paragraph()
            for topic in topics:
                p.add_run(f" • {topic} ")
            
            doc.add_heading('Simplified Content', level=1)
            doc.add_paragraph(simplified_text)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                doc.save(tmp_docx.name)
                print("☁️ Uploading Word to Cloudinary...")
                docx_upload = cloudinary.uploader.upload(
                    tmp_docx.name,
                    folder="eta-web-docs",
                    resource_type="raw"
                )
                docx_url = docx_upload.get("secure_url")
                docx_public_id = docx_upload.get("public_id")
                os.unlink(tmp_docx.name)
            
            print(f"✅ Web Word version uploaded: {docx_url}")
        except Exception as docx_err:
            print(f"⚠️ Word generation failed: {docx_err}")

        # Cleanup screenshot temp if exists
        if screenshot_temp_path and os.path.exists(screenshot_temp_path):
            os.unlink(screenshot_temp_path)

        return {
            "text": simplified_text,
            "raw_text": cleaned_text,
            "title": title,
            "url": url,
            "summary": summary,
            "topics": topics,
            "keywords": keywords,
            "thumbnail_url": thumbnail_url,
            "thumbnail_public_id": thumbnail_public_id,
            "pdf_url": pdf_url,
            "pdf_public_id": pdf_public_id,
            "docx_url": docx_url,
            "docx_public_id": docx_public_id,
            "metadata": {
                "extracted_at": "Now",
                "method": "WebScraping-LLM",
                "is_simplified": True
            }
        }

    except Exception as e:
        print(f"❌ Error in extract_web_content: {str(e)}")
        raise e
